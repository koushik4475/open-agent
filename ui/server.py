# openagent/ui/server.py
"""
Flask API server to connect the web UI to the OpenAgent backend.
Provides REST API endpoints for chat and system status.

Async model: one persistent background event loop (started at import)
runs ALL agent coroutines via asyncio.run_coroutine_threadsafe. This
replaces the old per-request asyncio.run(), which paid loop startup /
teardown on every call and tore down any pending async work.
"""

from __future__ import annotations

from flask import Flask, request, jsonify, send_from_directory, Response
from flask_cors import CORS
import asyncio
import sys
import os
import threading
from pathlib import Path

# Add parent directory of the project to path to support 'import openagent'
project_root = Path(__file__).resolve().parent.parent
sys.path.insert(0, str(project_root.parent))
os.chdir(project_root)

# Now import from the project
from openagent.core.agent import Agent, SYSTEM_PROMPT
from openagent.core.network import check_connectivity
from openagent.core.llm import LLMClient
from openagent.core.router import route

app = Flask(__name__, static_folder='.')
app.config['MAX_CONTENT_LENGTH'] = 36 * 1024 * 1024  # 36 MB max upload size
CORS(app)  # Enable CORS for local development

# Setup logging
import logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s',
    handlers=[
        logging.FileHandler("logs/openagent.log"),
        logging.StreamHandler()
    ]
)
logger = logging.getLogger("server")

# ── Persistent event loop ───────────────────────────────────────
# One background loop shared by every request. Flask worker threads
# submit coroutines with run_coroutine_threadsafe and block on the
# Future. Benefits: no per-request loop startup/teardown, and the
# singleton Agent's async resources always live on a single loop.
_async_loop = asyncio.new_event_loop()


def _loop_worker() -> None:
    asyncio.set_event_loop(_async_loop)
    _async_loop.run_forever()


threading.Thread(target=_loop_worker, name="openagent-eventloop", daemon=True).start()


def run_async(coro, timeout: float | None = 300):
    """Run a coroutine on the shared loop from a Flask worker thread."""
    return asyncio.run_coroutine_threadsafe(coro, _async_loop).result(timeout)


# Global agent instance
agent_instance = None
conversation_sessions = {}  # Map session_id -> list of messages
_agent_init_lock: asyncio.Lock | None = None  # Created lazily on the shared loop


async def get_agent():
    """Get or create the agent instance (single-flight).

    NOTE: this must use an asyncio.Lock, not a threading.Lock — all
    coroutines run on the ONE shared loop thread, so blocking on a
    threading lock while the holder is parked at an await would
    deadlock the loop.
    """
    global agent_instance, _agent_init_lock
    if agent_instance is not None:
        return agent_instance
    if _agent_init_lock is None:
        # No await between check and assignment → atomic on the single loop thread
        _agent_init_lock = asyncio.Lock()
    async with _agent_init_lock:
        if agent_instance is None:
            agent_instance = await Agent.create()
    return agent_instance


@app.route('/')
def index():
    """Serve the main UI page."""
    return send_from_directory('.', 'index.html')


@app.route('/api/upload', methods=['POST'])
def upload_file():
    """
    Handle file uploads.
    Saves the file to 'uploads' directory and returns the absolute path.
    """
    if 'file' not in request.files:
        return jsonify({'status': 'error', 'message': 'No file part'}), 400
    
    file = request.files['file']
    if file.filename == '':
        return jsonify({'status': 'error', 'message': 'No selected file'}), 400
        
    if file:
        upload_dir = Path(os.getcwd()) / "uploads"
        upload_dir.mkdir(exist_ok=True)
        
        # Safe filename
        filename = file.filename
        filepath = upload_dir / filename
        file.save(filepath)
        
        return jsonify({
            'status': 'success', 
            'filepath': str(filepath.absolute()),
            'filename': filename
        })

@app.route('/api/chat', methods=['POST'])
def chat():
    """
    Handle chat messages from the UI.
    
    Request JSON:
        {
            "message": "user message text"
        }
    
    Response JSON:
        {
            "response": "agent response text",
            "status": "success" | "error"
        }
    """
    global conversation_sessions
    
    try:
        data = request.get_json()
        user_message = data.get('message', '').strip()
        session_id = data.get('session_id', 'default')
        
        if not user_message:
            return jsonify({
                'status': 'error',
                'response': 'Empty message received.'
            }), 400
        
        # Initialize session if not exists
        if session_id not in conversation_sessions:
            conversation_sessions[session_id] = []
        
        # Run agent
        async def process_message():
            # Use global variable inside this async scope
            global conversation_sessions
            history = conversation_sessions[session_id]
            
            agent = await get_agent()
            response = await agent.run(user_message, history)
            
            # Update conversation history
            history.append({"role": "user", "content": user_message})
            history.append({"role": "assistant", "content": response})
            
            # Trim history to last 40 messages
            if len(history) > 40:
                conversation_sessions[session_id] = history[-40:]
            
            return response
        
        # Submit to the persistent background loop (no per-request loop churn)
        response = run_async(process_message())
        
        return jsonify({
            'status': 'success',
            'session_id': session_id,
            'response': response
        })
    
    except Exception as e:
        return jsonify({
            'status': 'error',
            'response': f'Error processing message: {str(e)}'
        }), 500


@app.route('/api/chat/stream', methods=['POST'])
def chat_stream():
    """
    SSE streaming chat endpoint — tokens arrive in real-time.
    Returns Server-Sent Events with each token chunk.
    """
    global conversation_sessions

    try:
        data = request.get_json()
        user_message = data.get('message', '').strip()
        session_id = data.get('session_id', 'default')

        if not user_message:
            return jsonify({'status': 'error', 'response': 'Empty message'}), 400

        if session_id not in conversation_sessions:
            conversation_sessions[session_id] = []

        def generate_stream():
            import json as _json
            history = conversation_sessions[session_id]

            # Async prep runs on the shared background loop; the token
            # loop itself (agent.llm.stream_generate) is a plain sync
            # generator that runs right here in the Flask thread.
            agent = run_async(get_agent())

            # Route the message
            tool_name, ctx = run_async(route(user_message))

            # For GENERAL chat, use streaming directly
            from openagent.core.router import ToolName
            if tool_name == ToolName.GENERAL:
                # Build prompt same way as agent does
                from openagent.core.agent import Agent as AgentClass
                memory_ctx = ""
                if not AgentClass._is_simple_query(user_message):
                    memory_ctx = run_async(agent.memory.retrieve(user_message))

                prompt = AgentClass._build_prompt(ctx["prompt"], memory_ctx)
                offline_warning = ctx.get("offline_warning", "")
                if offline_warning:
                    prompt = f"[NOTE: {offline_warning}]\n\n" + prompt

                full_response = []
                for token in agent.llm.stream_generate(prompt, system=SYSTEM_PROMPT, history=history):
                    full_response.append(token)
                    yield f"data: {_json.dumps({'token': token})}\n\n"

                complete = "".join(full_response)
            elif tool_name == ToolName.WEB_SEARCH:
                # For web search: do the search, then stream the LLM synthesis
                from openagent.tools.online.web_search import web_search
                from openagent.core.agent import Agent as AgentClass
                memory_ctx = ""
                if not AgentClass._is_simple_query(user_message):
                    memory_ctx = run_async(agent.memory.retrieve(user_message))
                search_results = run_async(web_search(ctx.get("query", ctx["prompt"])))
                prompt = AgentClass._build_prompt(ctx["prompt"], memory_ctx, web_results=search_results)

                full_response = []
                for token in agent.llm.stream_generate(prompt, system=SYSTEM_PROMPT, history=history):
                    full_response.append(token)
                    yield f"data: {_json.dumps({'token': token})}\n\n"

                complete = "".join(full_response)
            else:
                # For tool-based routes, run the full agent (non-streaming) and send as one chunk
                # (agent.run already queues its own background memory store)
                complete = run_async(agent.run(user_message, history))
                yield f"data: {_json.dumps({'token': complete})}\n\n"

            # Update session history
            history.append({"role": "user", "content": user_message})
            history.append({"role": "assistant", "content": complete})
            if len(history) > 40:
                conversation_sessions[session_id] = history[-40:]

            # Store in memory without delaying the final SSE event
            # (skip for tool routes — agent.run stored it already)
            if tool_name in (ToolName.GENERAL, ToolName.WEB_SEARCH):
                agent.memory.store_background(user_message, complete)

            yield f"data: {_json.dumps({'done': True})}\n\n"

        return Response(generate_stream(), mimetype='text/event-stream',
                        headers={'Cache-Control': 'no-cache', 'X-Accel-Buffering': 'no'})

    except Exception as e:
        import json as _json
        return Response(f"data: {_json.dumps({'error': str(e)})}\n\n",
                        mimetype='text/event-stream')


@app.route('/api/status', methods=['GET'])
def status():
    """
    Get system status information.
    
    Response JSON:
        {
            "llm_model": "phi3:mini",
            "llm_host": "http://localhost:11434",
            "network_online": true,
            "memory_db": "./data/chroma_db",
            "status": "success"
        }
    """
    try:
        async def get_status():
            agent = await get_agent()
            online = await check_connectivity()
            
            return {
                'status': 'success',
                'llm_model': agent.cfg.llm.model,
                'llm_host': agent.cfg.llm.host,
                'network_online': online,
                'memory_db': str(agent.cfg.memory.db_path)
            }
        
        status_data = run_async(get_status())

        return jsonify(status_data)
    
    except Exception as e:
        return jsonify({
            'status': 'error',
            'message': f'Error getting status: {str(e)}'
        }), 500


# Static tool registry for the UI — kept in sync with
# openagent.core.router.ToolName. (The Agent has no .tools attribute;
# the old implementation iterated agent.tools and always returned 500.)
TOOL_REGISTRY = [
    {"name": "general", "icon": "💬", "description": "General Q&A and conversation with the LLM", "category": "offline"},
    {"name": "summarize", "icon": "📝", "description": "Summarize text into key points and a TL;DR", "category": "offline"},
    {"name": "parse_file", "icon": "📄", "description": "Parse TXT, PDF and DOCX files and answer questions about them", "category": "offline"},
    {"name": "ocr_image", "icon": "🖼️", "description": "Extract text from images (OCR)", "category": "offline"},
    {"name": "analyze_image", "icon": "👁️", "description": "Vision AI image analysis (people, objects, scenes)", "category": "online"},
    {"name": "run_command", "icon": "🔧", "description": "Execute sandboxed shell commands", "category": "offline"},
    {"name": "file_ops", "icon": "📂", "description": "Read, write, search and fix project files (MCP)", "category": "offline"},
    {"name": "web_search", "icon": "🌐", "description": "Search the web via DuckDuckGo", "category": "online"},
    {"name": "web_fetch", "icon": "🔗", "description": "Fetch and read a webpage", "category": "online"},
]


@app.route('/api/tools', methods=['GET'])
def tools():
    """
    Get the list of available tools (static registry).

    Response JSON:
        {
            "status": "success",
            "tools": [
                {"name": str, "icon": str, "description": str,
                 "category": "offline" | "online", "available": bool}
            ]
        }

    `available` is True for offline tools; for online tools it reflects
    current connectivity (TTL-cached check, computed once per request).
    """
    try:
        online = run_async(check_connectivity())
    except Exception:
        online = False

    tool_list = [
        {**tool, "available": True if tool["category"] == "offline" else bool(online)}
        for tool in TOOL_REGISTRY
    ]
    return jsonify({
        'status': 'success',
        'tools': tool_list
    })


@app.route('/api/analyze-image', methods=['POST'])
def analyze_image():
    """
    Handle image upload and analysis via vision AI.
    
    Accepts a file upload, saves it, then runs the full
    vision analysis pipeline (vision AI → web search → synthesis).
    """
    global conversation_sessions
    session_id = request.form.get('session_id', 'default')

    if 'file' not in request.files:
        return jsonify({'status': 'error', 'response': 'No file uploaded'}), 400
    
    file = request.files['file']
    if file.filename == '':
        return jsonify({'status': 'error', 'response': 'No file selected'}), 400
    
    try:
        # Save the uploaded file
        upload_dir = Path(os.getcwd()) / "uploads"
        upload_dir.mkdir(exist_ok=True)
        filepath = upload_dir / file.filename
        file.save(filepath)
        
        # Run the analysis through the agent
        async def do_analysis():
            agent = await get_agent()
            history = conversation_sessions.setdefault(session_id, [])

            # Construct file tag for router
            user_msg = f"Analyse this file: [FILE:{filepath.absolute()}]"
            response = await agent.run(user_msg, history)

            # Update conversation history
            history.append({"role": "user", "content": f"[Uploaded image: {file.filename}]"})
            history.append({"role": "assistant", "content": response})

            if len(history) > 40:
                conversation_sessions[session_id] = history[-40:]

            return response
        
        response = run_async(do_analysis())

        return jsonify({
            'status': 'success',
            'response': response,
            'filename': file.filename
        })
    
    except Exception as e:
        return jsonify({
            'status': 'error',
            'response': f'Image analysis failed: {str(e)}'
        }), 500


@app.route('/api/settings', methods=['POST'])
def update_settings():
    """Update agent settings (project path for MCP)."""
    data = request.get_json()
    project_path = data.get('project_path', '')

    from openagent.tools.offline.file_ops import set_project_path
    set_project_path(project_path)

    return jsonify({'status': 'success', 'project_path': project_path})


@app.route('/api/export', methods=['POST'])
def export_response():
    """Export a response as PDF or DOCX with proper filenames."""
    data = request.get_json()
    text = data.get('text', '')
    fmt = data.get('format', 'txt')

    if not text:
        return jsonify({'status': 'error', 'message': 'No text to export'}), 400

    from flask import make_response
    import io

    if fmt == 'docx':
        try:
            from docx import Document
            doc = Document()
            doc.add_heading('OpenAgent Response', level=1)
            for paragraph in text.split('\n'):
                if paragraph.strip():
                    doc.add_paragraph(paragraph)
            buf = io.BytesIO()
            doc.save(buf)
            buf.seek(0)

            response = make_response(buf.read())
            response.headers['Content-Type'] = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            response.headers['Content-Disposition'] = 'attachment; filename=openagent_response.docx'
            return response
        except ImportError:
            response = make_response(text.encode('utf-8'))
            response.headers['Content-Type'] = 'text/plain; charset=utf-8'
            response.headers['Content-Disposition'] = 'attachment; filename=openagent_response.txt'
            return response

    elif fmt == 'pdf':
        try:
            from reportlab.lib.pagesizes import letter
            from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
            from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
            from reportlab.lib.units import inch

            # Register a Unicode font (Latin + Kannada) so non-Latin text
            # renders instead of empty boxes. Load via an absolute path
            # derived from this file's location so it works regardless of cwd.
            title_style = None
            body_style = None
            try:
                from reportlab.pdfbase import pdfmetrics
                from reportlab.pdfbase.ttfonts import TTFont as RLTTFont

                _FONT_PATH = Path(__file__).resolve().parent / "fonts" / "NotoSansKannada.ttf"
                if "NotoKannada" not in pdfmetrics.getRegisteredFontNames():
                    pdfmetrics.registerFont(RLTTFont("NotoKannada", str(_FONT_PATH)))

                base_styles = getSampleStyleSheet()
                title_style = ParagraphStyle(
                    'NotoTitle', parent=base_styles['Title'], fontName="NotoKannada")
                body_style = ParagraphStyle(
                    'NotoBody', parent=base_styles['Normal'], fontName="NotoKannada")
            except Exception:
                # Font registration failed for any reason: fall back to the
                # default styles so the PDF still generates (Latin only).
                title_style = None
                body_style = None

            buf = io.BytesIO()
            doc = SimpleDocTemplate(buf, pagesize=letter)
            styles = getSampleStyleSheet()
            if title_style is None:
                title_style = styles['Title']
            if body_style is None:
                body_style = styles['Normal']
            story = []
            story.append(Paragraph('OpenAgent Response', title_style))
            story.append(Spacer(1, 0.2 * inch))

            for line in text.split('\n'):
                if line.strip():
                    safe = line.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
                    story.append(Paragraph(safe, body_style))
                else:
                    story.append(Spacer(1, 0.1 * inch))

            doc.build(story)
            buf.seek(0)

            response = make_response(buf.read())
            response.headers['Content-Type'] = 'application/pdf'
            response.headers['Content-Disposition'] = 'attachment; filename=openagent_response.pdf'
            return response
        except ImportError:
            response = make_response(text.encode('utf-8'))
            response.headers['Content-Type'] = 'text/plain; charset=utf-8'
            response.headers['Content-Disposition'] = 'attachment; filename=openagent_response.txt'
            return response

    else:
        response = make_response(text.encode('utf-8'))
        response.headers['Content-Type'] = 'text/plain; charset=utf-8'
        response.headers['Content-Disposition'] = 'attachment; filename=openagent_response.txt'
        return response


@app.route('/api/clear', methods=['POST'])
def clear_history():
    """Clear conversation history for a session."""
    global conversation_sessions
    data = request.get_json(silent=True) or {}
    session_id = data.get('session_id', 'default')
    conversation_sessions[session_id] = []
    return jsonify({
        'status': 'success',
        'message': 'Conversation cleared.'
    })


if __name__ == "__main__":
    import os
    port = int(os.environ.get("PORT", 7860))
    app.run(host="0.0.0.0", port=port, debug=False)
