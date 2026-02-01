# openagent/tests/test_parsers.py
"""
Tests for all file parsers.
Uses temporary files — no external dependencies.
"""

import pytest
import tempfile
from pathlib import Path

from openagent.parsers import txt_parser, unified


# ─── TXT Parser ────────────────────────────────────────────────

class TestTxtParser:
    def test_basic_read(self, tmp_path):
        f = tmp_path / "hello.txt"
        f.write_text("Hello, World!", encoding="utf-8")
        assert txt_parser.extract_text(f) == "Hello, World!"

    def test_utf8_characters(self, tmp_path):
        f = tmp_path / "unicode.txt"
        f.write_text("Héllo Wörld — café ☕", encoding="utf-8")
        assert "café" in txt_parser.extract_text(f)

    def test_latin1_fallback(self, tmp_path):
        f = tmp_path / "latin.txt"
        f.write_bytes("caf\xe9".encode("latin-1"))
        result = txt_parser.extract_text(f)
        assert "caf" in result

    def test_file_not_found(self):
        with pytest.raises(FileNotFoundError):
            txt_parser.extract_text(Path("/nonexistent/file.txt"))

    def test_empty_file(self, tmp_path):
        f = tmp_path / "empty.txt"
        f.write_text("", encoding="utf-8")
        assert txt_parser.extract_text(f) == ""


# ─── Unified Parser ───────────────────────────────────────────

class TestUnifiedParser:
    def test_dispatches_txt(self, tmp_path):
        f = tmp_path / "test.txt"
        f.write_text("unified test content", encoding="utf-8")
        result = unified.parse_file(f)
        assert result == "unified test content"

    def test_unsupported_extension(self, tmp_path):
        f = tmp_path / "test.xyz"
        f.write_text("data", encoding="utf-8")
        with pytest.raises(ValueError, match="Unsupported file format"):
            unified.parse_file(f)

    def test_file_not_found(self):
        with pytest.raises(FileNotFoundError):
            unified.parse_file(Path("/no/such/file.txt"))

    def test_supported_extensions_list(self):
        exts = unified.supported_extensions()
        assert ".txt" in exts
        assert ".pdf" in exts
        assert ".docx" in exts
        assert ".png" in exts


# ─── DOCX Parser (requires python-docx) ───────────────────────

class TestDocxParser:
    def test_basic_docx(self, tmp_path):
        """Create a minimal DOCX and parse it."""
        from docx import Document
        from openagent.parsers import docx_parser

        doc = Document()
        doc.add_paragraph("First paragraph.")
        doc.add_paragraph("Second paragraph.")

        f = tmp_path / "test.docx"
        doc.save(str(f))

        result = docx_parser.extract_text(f)
        assert "First paragraph." in result
        assert "Second paragraph." in result

    def test_docx_with_table(self, tmp_path):
        from docx import Document
        from openagent.parsers import docx_parser

        doc = Document()
        table = doc.add_table(rows=2, cols=2)
        table.cell(0, 0).text = "Name"
        table.cell(0, 1).text = "Age"
        table.cell(1, 0).text = "Alice"
        table.cell(1, 1).text = "30"

        f = tmp_path / "table.docx"
        doc.save(str(f))

        result = docx_parser.extract_text(f)
        assert "Alice" in result
        assert "30" in result

    def test_docx_not_found(self):
        from openagent.parsers import docx_parser
        with pytest.raises(FileNotFoundError):
            docx_parser.extract_text(Path("/nope.docx"))


# ─── PDF Parser (requires PyMuPDF) ────────────────────────────

class TestPdfParser:
    def test_basic_pdf(self, tmp_path):
        """Create a minimal PDF with PyMuPDF and parse it back."""
        import fitz
        from openagent.parsers import pdf_parser

        doc = fitz.open()
        page = doc.new_page(width=612, height=792)
        page.insert_text((72, 72), "Hello from a test PDF.", fontsize=12)

        f = tmp_path / "test.pdf"
        doc.save(str(f))
        doc.close()

        result = pdf_parser.extract_text(f)
        assert "Hello from a test PDF." in result

    def test_pdf_not_found(self):
        from openagent.parsers import pdf_parser
        with pytest.raises(FileNotFoundError):
            pdf_parser.extract_text(Path("/nope.pdf"))
