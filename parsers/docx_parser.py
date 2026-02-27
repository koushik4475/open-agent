# openagent/parsers/docx_parser.py
"""
DOCX → Text extraction using python-docx.

Why python-docx:
  - MIT license.
  - Handles paragraphs, tables, headers, footers.
  - Pure Python. No external binary. Offline.
"""

from __future__ import annotations
from pathlib import Path
import logging

from docx import Document

logger = logging.getLogger("openagent.parsers.docx")


def extract_text(filepath: Path) -> str:
    """
    Extract text from a .docx file.
    Includes paragraphs and table cell contents.
    """
    if not filepath.exists():
        raise FileNotFoundError(f"DOCX not found: {filepath}")

    try:
        doc = Document(str(filepath))
        sections: list[str] = []

        # ── Paragraphs ────────────────────────────────────────
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                # Preserve heading structure (guard against None style)
                style_name = para.style.name if para.style else ""
                if style_name.startswith("Heading"):
                    sections.append(f"\n## {text}")
                else:
                    sections.append(text)

        # ── Tables ────────────────────────────────────────────
        for table_idx, table in enumerate(doc.tables, 1):
            sections.append(f"\n[Table {table_idx}]")
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                sections.append(" | ".join(cells))

        full_text = "\n".join(sections)

        if not full_text.strip():
            return "[DOCX file is empty or contains only formatting.]"

        logger.info(f"Extracted {len(full_text)} chars from {filepath.name}")
        return full_text

    except Exception as e:
        logger.error(f"DOCX extraction error: {e}")
        raise RuntimeError(f"Failed to parse DOCX '{filepath.name}': {e}")
